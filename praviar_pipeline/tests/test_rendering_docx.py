"""Tests for praviar_pipeline.rendering.docx_report -- DOCX export for FTO reports."""

from __future__ import annotations

import base64
import io
from datetime import UTC, date, datetime

import pytest

from praviar_pipeline.models.analysis import (
    ClaimAnalysis,
    ClaimElement,
    ElementStatus,
    PatentAnalysis,
    RiskLevel,
)
from praviar_pipeline.models.audit import (
    PipelineAuditTrail,
    SearchFunnelEntry,
    StepTiming,
)
from praviar_pipeline.models.compound import ResolvedCompound
from praviar_pipeline.models.report import (
    FTOReport,
    RiskSummary,
    SourceHealth,
    SourceHealthEntry,
    SourceStatus,
)
from praviar_pipeline.models.verification import VerificationCheck, VerificationResult
from praviar_pipeline.rendering.branding import BrandingConfig
from praviar_pipeline.rendering.docx_report import render_docx
from praviar_pipeline.rendering.export_options import ExportRenderOptions

docx = pytest.importorskip("docx")

_PNG_1X1 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def minimal_compound() -> ResolvedCompound:
    return ResolvedCompound(
        name="Aspirin",
        canonical_smiles="CC(=O)Oc1ccccc1C(=O)O",
        inchi="InChI=1S/C9H8O4/c1-6(10)13-8-5-3-2-4-7(8)9(11)12/h2-5H,1H3,(H,11,12)",
        inchi_key="BSYNRYMUTXBXSQ-UHFFFAOYSA-N",
        molecular_formula="C9H8O4",
        molecular_weight=180.16,
        pubchem_cid=2244,
        original_input="aspirin",
        input_type="name",
    )


@pytest.fixture
def minimal_analysis() -> PatentAnalysis:
    return PatentAnalysis(
        patent_id="US7851188B2",
        title="Methods for producing aspirin derivatives",
        assignee="PharmaCorp Inc.",
        expiry_date=date(2028, 3, 15),
        claims_analyzed=[
            ClaimAnalysis(
                claim_number=1,
                claim_type="independent",
                elements=[
                    ClaimElement(
                        element_number=1,
                        element_text="A method for producing acetylsalicylic acid",
                        status=ElementStatus.MET,
                        reasoning="Target compound is acetylsalicylic acid (aspirin)",
                        confidence=0.95,
                    ),
                    ClaimElement(
                        element_number=2,
                        element_text="comprising reacting salicylic acid with acetic anhydride",
                        status=ElementStatus.NOT_MET,
                        reasoning="Different synthesis route used",
                        confidence=0.9,
                    ),
                ],
                overall_status=ElementStatus.NOT_MET,
                overall_confidence=0.9,
            ),
        ],
        risk_level=RiskLevel.MEDIUM,
        risk_summary="Moderate risk -- one claim element not met",
    )


@pytest.fixture
def high_risk_analysis() -> PatentAnalysis:
    """A HIGH risk analysis where all elements are MET."""
    return PatentAnalysis(
        patent_id="US8888888B2",
        title="Aspirin formulation process",
        assignee="GreenChem Corp",
        expiry_date=date(2035, 1, 1),
        claims_analyzed=[
            ClaimAnalysis(
                claim_number=1,
                claim_type="independent",
                elements=[
                    ClaimElement(
                        element_number=1,
                        element_text="producing acetylsalicylic acid",
                        status=ElementStatus.MET,
                        reasoning="Exact match",
                        confidence=0.99,
                    ),
                ],
                overall_status=ElementStatus.MET,
                overall_confidence=0.99,
            ),
        ],
        risk_level=RiskLevel.HIGH,
        risk_summary="All claim elements met -- direct infringement risk",
    )


@pytest.fixture
def rendering_report(
    minimal_compound: ResolvedCompound,
    minimal_analysis: PatentAnalysis,
    high_risk_analysis: PatentAnalysis,
) -> FTOReport:
    """FTOReport with all fields populated for rendering tests."""
    now = datetime.now(UTC)
    return FTOReport(
        report_id="docx-test-001",
        generated_at=now,
        compound=minimal_compound,
        risk_summary=RiskSummary(
            overall_risk=RiskLevel.HIGH,
            blocking_patents_count=2,
            total_patents_analyzed=5,
            key_risks=[
                "US7851188B2: medium risk",
                "US8888888B2: blocks composition",
            ],
            executive_summary="Analysis reveals 2 blocking patents for Aspirin.",
        ),
        patent_analyses=[minimal_analysis, high_risk_analysis],
        verification=VerificationResult(
            checks=[
                VerificationCheck(
                    check_name="citation_grounding",
                    passed=True,
                    details="All patent IDs found in search results",
                ),
            ],
            all_citations_valid=True,
            all_claims_grounded=True,
            all_entities_valid=True,
            dates_consistent=True,
            risk_levels_justified=True,
        ),
        total_patents_found=50,
        patents_after_triage=2,
        search_sources_used=["pubchem", "bigquery"],
        source_health=SourceHealth(
            entries=[
                SourceHealthEntry(
                    source="pubchem",
                    status=SourceStatus.OK,
                    patent_count=42,
                ),
                SourceHealthEntry(
                    source="bigquery",
                    status=SourceStatus.FAILED,
                    patent_count=0,
                    error_message="HTTP 504 timeout from BigQuery",
                ),
                SourceHealthEntry(
                    source="lens",
                    status=SourceStatus.NOT_CONFIGURED,
                    patent_count=0,
                    error_message="API key missing",
                ),
            ]
        ),
        audit_trail=PipelineAuditTrail(
            search_funnel=[
                SearchFunnelEntry(patent_id="US7851188B2", included_in_triage=True),
            ],
            timing_data=[
                StepTiming(
                    step_name="step2_search",
                    started_at=now,
                    completed_at=now,
                    duration_seconds=3.5,
                    items_processed=100,
                    items_output=50,
                ),
            ],
            total_patents_discovered=100,
        ),
        patent_narratives={
            "US7851188B2": "This patent covers aspirin derivative production methods.",
        },
        llm_models_used={
            "triage": "claude-haiku-4-5-20251001",
            "analysis": "claude-sonnet-4-20250514",
        },
    )


@pytest.fixture
def empty_analyses_report(minimal_compound: ResolvedCompound) -> FTOReport:
    """FTOReport with no patent_analyses (edge case)."""
    return FTOReport(
        report_id="docx-test-empty",
        compound=minimal_compound,
        risk_summary=RiskSummary(
            overall_risk=RiskLevel.CLEAR,
            blocking_patents_count=0,
            total_patents_analyzed=0,
            key_risks=[],
            executive_summary="No blocking patents identified.",
        ),
        patent_analyses=[],
        total_patents_found=0,
        patents_after_triage=0,
        search_sources_used=["pubchem"],
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _full_text(document) -> str:
    """Extract all paragraph text from a python-docx Document."""
    parts = []
    for para in document.paragraphs:
        parts.append(para.text)
    # Also extract text from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _write_png_logo(path) -> None:
    path.write_bytes(base64.b64decode(_PNG_1X1))


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestDocxRendering:
    """Tests for the DOCX report renderer."""

    def test_returns_bytes(self, rendering_report: FTOReport):
        """render_docx should return bytes."""
        result = render_docx(rendering_report)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_zip_magic_bytes(self, rendering_report: FTOReport):
        """DOCX files are ZIP archives starting with PK header."""
        result = render_docx(rendering_report)
        assert result[:4] == b"PK\x03\x04"

    def test_loadable_as_document(self, rendering_report: FTOReport):
        """The returned bytes should load as a valid python-docx Document."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_executive_verdict_uses_clearance_decision(self, rendering_report: FTOReport):
        from docx import Document

        rendering_report.risk_summary.overall_risk = RiskLevel.CLEAR
        result = render_docx(rendering_report)
        text = _full_text(Document(io.BytesIO(result)))

        assert "Clearance Decision: UNCLEAR" in text
        assert "Overall Risk: Clear" not in text

    def test_executive_scope_omits_detailed_sections(self, rendering_report: FTOReport):
        """Executive-only DOCX exports keep caveats but omit detailed bodies."""
        from docx import Document

        options = ExportRenderOptions.from_values(
            ["executive_summary"],
            audience="executive",
        )
        result = render_docx(rendering_report, options=options)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)

        assert "Evidence Scope" in text
        assert "Export audience: Executive Brief" in text
        assert "Executive Summary" in text
        assert "Strategic Recommendations" in text
        assert "Risk Assessment Matrix" not in text
        assert "Detailed Patent Analysis" not in text
        assert "Invalidity Screening" not in text

    def test_preserves_long_claim_chart_text(self, rendering_report: FTOReport):
        """Claim chart DOCX tables should preserve full legal text."""
        from docx import Document

        claim = rendering_report.patent_analyses[0].claims_analyzed[0]
        tail = "SENTINEL-REPORT-CLAIM-TAIL-DO-NOT-TRUNCATE"
        reasoning_tail = "SENTINEL-REPORT-REASONING-TAIL-DO-NOT-TRUNCATE"
        claim.elements[0].element_text = "claim text " + ("alpha " * 90) + tail
        claim.elements[0].reasoning = "reasoning " + ("beta " * 120) + reasoning_tail

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)

        assert tail in text
        assert reasoning_tail in text

    def test_title_present(self, rendering_report: FTOReport):
        """Document should contain the FTO analysis title."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "Freedom-to-Operate" in text

    def test_compound_name_present(self, rendering_report: FTOReport):
        """Document should contain the compound name."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "Aspirin" in text

    def test_risk_level_present(self, rendering_report: FTOReport):
        """Document should contain risk level text."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "HIGH" in text

    def test_patent_ids_present(self, rendering_report: FTOReport):
        """Document should contain patent IDs from analyses."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "US7851188B2" in text
        assert "US8888888B2" in text

    def test_tables_exist(self, rendering_report: FTOReport):
        """Document should contain tables (risk matrix, compound profile, etc.)."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) > 0

    def test_risk_matrix_table(self, rendering_report: FTOReport):
        """Risk matrix table should have patent data rows."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))

        # Find the risk matrix table by looking for Patent ID in header
        found_risk_matrix = False
        for table in doc.tables:
            header_texts = [cell.text for cell in table.rows[0].cells]
            if "Patent ID" in header_texts and "Risk" in header_texts:
                found_risk_matrix = True
                # Should have at least header + data rows
                assert len(table.rows) > 1
                break
        assert found_risk_matrix, "Risk matrix table not found in document"

    def test_executive_summary_present(self, rendering_report: FTOReport):
        """Executive summary text should appear in the document."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "Clearance decision: UNCLEAR." in text
        assert "Analysis reveals 2 blocking patents" not in text

    def test_evidence_scope_precedes_executive_risk(
        self,
        rendering_report: FTOReport,
    ):
        """DOCX front matter should show source scope before risk conclusions."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)

        assert "1. Evidence Scope" in text
        assert "1 of 3 configured sources completed" in text
        assert "Review required before relying on absence-of-risk conclusions" in text
        assert "bigquery" in text
        assert "lens" in text
        assert text.index("1. Evidence Scope") < text.index("2. Executive Summary")

    def test_claim_chart_section(self, rendering_report: FTOReport):
        """Claim chart tables should be present when claims exist."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        # Claim element text should appear
        assert "Claim 1" in text

    def test_claim_element_statuses(self, rendering_report: FTOReport):
        """Claim element MET/NOT MET statuses should appear in tables."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "MET" in text
        assert "NOT MET" in text

    def test_unresolved_doe_and_fwr_are_not_rendered_as_negative(
        self,
        rendering_report: FTOReport,
        sample_doe_assessment,
    ):
        from docx import Document

        sample_doe_assessment.overall_equivalent = None
        sample_doe_assessment.fwr.same_function = None
        sample_doe_assessment.fwr.same_way = None
        sample_doe_assessment.fwr.same_result = None
        sample_doe_assessment.fwr.equivalent = None
        rendering_report.doe_assessments = [sample_doe_assessment]

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)

        assert "UNRESOLVED" in text
        assert "NOT EQUIVALENT" not in text
        assert "Function: Unresolved" in text

    def test_assignee_present(self, rendering_report: FTOReport):
        """Patent assignee names should appear in the document."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "PharmaCorp Inc." in text

    def test_report_id_present(self, rendering_report: FTOReport):
        """Report ID should appear in the document."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "docx-test-001" in text

    def test_disclaimer_present(self, rendering_report: FTOReport):
        """Disclaimer section should appear in the document."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "NOT constitute" in text or "does NOT constitute" in text


class TestDocxBranding:
    """Tests for DOCX rendering with BrandingConfig."""

    def test_with_privilege_header(self, rendering_report: FTOReport):
        """Organization branding cannot self-assert attorney privilege."""
        from docx import Document

        branding = BrandingConfig(
            privilege_header="PRIVILEGED AND CONFIDENTIAL -- ATTORNEY WORK PRODUCT",
        )
        result = render_docx(rendering_report, branding=branding)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "PRIVILEGED AND CONFIDENTIAL" not in text
        assert "CONFIDENTIAL DRAFT" in text

    def test_with_firm_name(self, rendering_report: FTOReport):
        """Firm name should appear in the document when set."""
        from docx import Document

        branding = BrandingConfig(firm_name="Baker & McKenzie LLP")
        result = render_docx(rendering_report, branding=branding)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "Baker & McKenzie LLP" in text

    def test_with_custom_disclaimer(self, rendering_report: FTOReport):
        """Custom disclaimer text appends without replacing the mandatory text."""
        from docx import Document

        custom_text = "This is a custom disclaimer for the firm."
        branding = BrandingConfig(disclaimer_text=custom_text)
        result = render_docx(rendering_report, branding=branding)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert custom_text in text
        assert "does NOT constitute legal advice" in text

    def test_with_matter_number(self, rendering_report: FTOReport):
        """Matter number should appear on cover page when set."""
        from docx import Document

        branding = BrandingConfig(matter_number="MATTER-2026-001")
        result = render_docx(rendering_report, branding=branding)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "MATTER-2026-001" in text

    def test_default_branding(self, rendering_report: FTOReport):
        """Rendering with no branding (default) should still produce valid DOCX."""
        result = render_docx(rendering_report, branding=None)
        assert isinstance(result, bytes)
        assert result[:4] == b"PK\x03\x04"

    def test_default_branding_includes_praviar_mark(self, rendering_report: FTOReport):
        """Default DOCX cover should include the Praviar mark image."""
        from docx import Document

        result = render_docx(rendering_report, branding=None)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "PRAVIAR" in text
        assert len(doc.inline_shapes) >= 1

    def test_custom_png_logo_uses_firm_identity_without_praviar_footer(
        self,
        rendering_report: FTOReport,
        tmp_path,
    ):
        """Custom logos should render and suppress generated-by-Praviar copy."""
        from docx import Document

        logo = tmp_path / "firm-logo.png"
        _write_png_logo(logo)
        branding = BrandingConfig(logo_path=str(logo), firm_name="Acme Counsel")

        result = render_docx(rendering_report, branding=branding)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)

        assert "Acme Counsel" in text
        assert "Generated by Praviar" not in text
        assert len(doc.inline_shapes) >= 1

    def test_custom_svg_logo_fails_with_clear_docx_error(
        self,
        rendering_report: FTOReport,
        tmp_path,
    ):
        """DOCX should reject unsupported SVG logos before python-docx crashes."""
        logo = tmp_path / "firm-logo.svg"
        logo.write_text("<svg xmlns='http://www.w3.org/2000/svg' />", encoding="utf-8")
        branding = BrandingConfig(logo_path=str(logo), firm_name="Acme Counsel")

        with pytest.raises(
            RuntimeError,
            match=r"DOCX does not support branding logo format '\.svg'",
        ):
            render_docx(rendering_report, branding=branding)

    def test_missing_custom_logo_fails_closed(
        self,
        rendering_report: FTOReport,
        tmp_path,
    ):
        """Missing logos should not silently produce an unbranded document."""
        branding = BrandingConfig(
            logo_path=str(tmp_path / "missing-logo.png"),
            firm_name="Acme Counsel",
        )

        with pytest.raises(RuntimeError, match="Branding logo not found for DOCX"):
            render_docx(rendering_report, branding=branding)


class TestDocxEdgeCases:
    """Tests for edge cases in DOCX rendering."""

    def test_empty_patent_analyses(self, empty_analyses_report: FTOReport):
        """Should produce valid DOCX even with no patent analyses."""
        result = render_docx(empty_analyses_report)
        assert isinstance(result, bytes)
        assert result[:4] == b"PK\x03\x04"

    def test_empty_analyses_loadable(self, empty_analyses_report: FTOReport):
        """Empty analyses report should be loadable as a Document."""
        from docx import Document

        result = render_docx(empty_analyses_report)
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_empty_analyses_no_risk_matrix_rows(self, empty_analyses_report: FTOReport):
        """With no analyses, the risk matrix section should not create data rows."""
        from docx import Document

        result = render_docx(empty_analyses_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        # The report should still render but US patent IDs should not appear
        assert "US7851188B2" not in text

    def test_compound_profile_table(self, rendering_report: FTOReport):
        """Compound profile should include SMILES, InChIKey, and molecular details."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "CC(=O)Oc1ccccc1C(=O)O" in text
        assert "BSYNRYMUTXBXSQ-UHFFFAOYSA-N" in text
        assert "C9H8O4" in text
        assert "180.16" in text

    def test_search_sources_present(self, rendering_report: FTOReport):
        """Configured source telemetry should appear in the methodology section."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "Configured source requests" in text
        assert "pubchem" in text
        assert "bigquery" in text
        assert "lens" in text
        assert "Source status" in text

    def test_provider_query_credentials_are_not_rendered(self, rendering_report: FTOReport):
        from docx import Document

        rendering_report.source_health.entries[
            1
        ].error_message = "401 https://api.openalex.org/works?api_key=SUPERSECRET"

        text = _full_text(Document(io.BytesIO(render_docx(rendering_report))))

        assert "Provider request failed" in text
        assert "SUPERSECRET" not in text
        assert "api_key=" not in text

    def test_verification_section_present(self, rendering_report: FTOReport):
        """Verification section should include check results."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "citation_grounding" in text
        assert "PASS" in text

    def test_patent_narrative_included(self, rendering_report: FTOReport):
        """Patent narratives should appear in the detailed analysis."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "aspirin derivative production methods" in text

    def test_llm_models_present(self, rendering_report: FTOReport):
        """LLM model attribution should appear in the document."""
        from docx import Document

        result = render_docx(rendering_report)
        doc = Document(io.BytesIO(result))
        text = _full_text(doc)
        assert "claude-haiku-4-5-20251001" in text
        assert "claude-sonnet-4-20250514" in text
