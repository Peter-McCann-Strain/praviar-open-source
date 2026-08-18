"""Focused tests for praviar_pipeline.rendering.docx_report_sections helpers."""

from __future__ import annotations

import pytest

from praviar_pipeline.rendering.branding import BrandingConfig
from praviar_pipeline.rendering.docx_report_sections import (
    add_cover_page,
    add_disclaimer,
    build_report_sections,
)

pytest.importorskip("docx")


def _document_text(document) -> str:
    parts: list[str] = []
    for para in document.paragraphs:
        parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_public_surface_keeps_build_report_sections(sample_report) -> None:
    from docx import Document

    doc = Document()
    build_report_sections(doc, sample_report, BrandingConfig())
    text = _document_text(doc)
    assert "Executive Summary" in text
    assert "Verification & Quality" in text
    assert "Appendices" in text


def test_cover_and_disclaimer_helpers_render_branding(sample_report) -> None:
    from docx import Document

    branding = BrandingConfig(
        privilege_header="PRIVILEGED AND CONFIDENTIAL",
        matter_number="MAT-123",
        disclaimer_text="Custom disclaimer text",
        firm_name="Test Firm",
        hide_praviar_pipeline_branding=True,
    )

    doc = Document()
    add_cover_page(doc, sample_report, branding)
    add_disclaimer(doc, sample_report, branding)
    text = _document_text(doc)
    assert "Freedom-to-Operate Analysis" in text
    assert "MAT-123" in text
    assert "Custom disclaimer text" in text
    assert "does NOT constitute legal advice" in text
    assert "PRIVILEGED AND CONFIDENTIAL" not in text
    assert "CONFIDENTIAL DRAFT" in text
    assert "Test Firm" in text
