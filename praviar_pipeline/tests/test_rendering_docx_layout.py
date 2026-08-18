"""Tests for DOCX layout helpers."""

from __future__ import annotations

from docx import Document

from praviar_pipeline.models.analysis import RiskLevel
from praviar_pipeline.rendering.branding import BrandingConfig
from praviar_pipeline.rendering.docx_report_layout import (
    configure_document,
    get_risk_level,
    status_bg,
)


def test_status_bg_mapping() -> None:
    assert status_bg("met") == "#FDECEC"
    assert status_bg("partially_met") == "#F7EEE5"
    assert status_bg("not_met") == "#D7ECE5"
    assert status_bg("unclear") == "#F6F4EF"
    assert status_bg("unknown") == "#F6F4EF"


def test_get_risk_level_mapping() -> None:
    assert get_risk_level("high") is RiskLevel.HIGH
    assert get_risk_level("medium") is RiskLevel.MEDIUM
    assert get_risk_level("low") is RiskLevel.LOW
    assert get_risk_level("clear") is RiskLevel.CLEAR
    assert get_risk_level("unknown") is RiskLevel.MEDIUM


def test_configure_document_applies_branding() -> None:
    doc = Document()
    branding = BrandingConfig(
        privilege_header="Header",
        confidentiality_footer="Footer",
    )

    configure_document(doc, branding)

    assert doc.styles["Normal"].font.name
    assert doc.sections[0].different_first_page_header_footer is True
    assert doc.sections[0].header.paragraphs[0].text == (
        "CONFIDENTIAL DRAFT | Praviar FTO Analysis"
    )
    assert "Header" not in doc.sections[0].header.paragraphs[0].text
    assert doc.sections[0].first_page_header.paragraphs[0].text == ""
    assert doc.sections[0].footer.paragraphs[0].text.startswith("Footer | Page ")
    assert "<w:drawing" in doc.sections[0].header.paragraphs[0]._element.xml


def test_configure_document_adds_footer_page_number_fields() -> None:
    doc = Document()
    branding = BrandingConfig(confidentiality_footer="Footer")

    configure_document(doc, branding)

    footer_xml = doc.sections[0].footer.paragraphs[0]._element.xml
    assert " PAGE " in footer_xml
    assert " NUMPAGES " in footer_xml


def test_configure_document_suppresses_praviar_mark_in_white_label_mode() -> None:
    doc = Document()
    branding = BrandingConfig(
        firm_name="Acme Counsel",
        hide_praviar_pipeline_branding=True,
    )

    configure_document(doc, branding)

    header = doc.sections[0].header.paragraphs[0]
    assert header.text == "CONFIDENTIAL DRAFT | Acme Counsel FTO Analysis"
    assert "<w:drawing" not in header._element.xml
