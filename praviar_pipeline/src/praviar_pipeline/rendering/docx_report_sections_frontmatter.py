"""Front matter sections for DOCX report rendering."""

from __future__ import annotations

from typing import TYPE_CHECKING

from praviar_pipeline.rendering.brand_mark import render_praviar_mark_png_stream
from praviar_pipeline.rendering.branding import (
    SUPPORTED_OFFICE_BRANDING_LOGO_EXTENSIONS,
    resolve_branding_logo_path,
)
from praviar_pipeline.rendering.design import (
    BRAND_DEEP_TEAL,
    BRAND_PAPER,
    BRAND_SECONDARY_TEXT,
    DOCX_FONTS,
    RISK_BG,
    RISK_FILL,
    risk_display,
    risk_sort_key,
)
from praviar_pipeline.rendering.docx_report_layout import (
    add_styled_paragraph,
    get_risk_level,
    set_cell_shading,
    style_body_cell,
    style_header_cell,
)
from praviar_pipeline.rendering.docx_report_layout import (
    rgb_color as _rgb_color,
)
from praviar_pipeline.rendering.evidence_scope import (
    format_source_list,
    source_status_detail,
    source_status_label,
    summarize_evidence_scope,
)
from praviar_pipeline.rendering.governed_decision import (
    governed_blocking_count,
    governed_decision_label,
    governed_executive_summary,
    governed_risk_level,
)

if TYPE_CHECKING:
    from praviar_pipeline.models.report import FTOReport
    from praviar_pipeline.rendering.branding import BrandingConfig


def add_cover_page(doc, report: FTOReport, branding: BrandingConfig) -> None:
    """Add cover page with title, compound, date, and confidentiality."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    logo_path = resolve_branding_logo_path(
        branding,
        renderer_name="DOCX",
        supported_extensions=SUPPORTED_OFFICE_BRANDING_LOGO_EXTENSIONS,
    )
    if logo_path is not None or not branding.suppresses_praviar_branding or branding.firm_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo_path is not None:
            p.add_run().add_picture(str(logo_path), width=Inches(1.0))
        elif not branding.suppresses_praviar_branding:
            p.add_run().add_picture(
                render_praviar_mark_png_stream(variant="on_light", size_px=384),
                width=Inches(1.0),
            )
        elif branding.firm_name:
            run = p.add_run(branding.firm_name)
            run.font.name = DOCX_FONTS["heading"]
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = _rgb_color(branding.primary_color)

        if not branding.suppresses_praviar_branding or branding.firm_name:
            brand_label = branding.display_name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(brand_label.upper())
            run.font.name = DOCX_FONTS["heading"]
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = _rgb_color(BRAND_SECONDARY_TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Freedom-to-Operate Analysis")
    run.font.name = DOCX_FONTS["heading"]
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = _rgb_color(branding.primary_color)
    p.paragraph_format.space_before = Pt(32)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(report.compound.name)
    run.font.name = DOCX_FONTS["heading"]
    run.font.size = Pt(18)
    run.font.color.rgb = _rgb_color(BRAND_DEEP_TEAL)
    p.paragraph_format.space_before = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M UTC')}")
    run.font.name = DOCX_FONTS["body"]
    run.font.size = Pt(11)
    run.font.color.rgb = _rgb_color(BRAND_SECONDARY_TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Report ID: {report.report_id}")
    run.font.name = DOCX_FONTS["code"]
    run.font.size = Pt(9)
    run.font.color.rgb = _rgb_color(BRAND_SECONDARY_TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL — NOT LEGAL ADVICE")
    run.font.name = DOCX_FONTS["heading"]
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = _rgb_color(RISK_FILL[get_risk_level("high")])
    p.paragraph_format.space_before = Pt(36)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(branding.legal_marking)
    run.font.name = DOCX_FONTS["heading"]
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = _rgb_color(BRAND_SECONDARY_TEXT)

    if branding.matter_number:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Matter: {branding.matter_number}")
        run.font.name = DOCX_FONTS["body"]
        run.font.size = Pt(10)
        run.font.color.rgb = _rgb_color(BRAND_SECONDARY_TEXT)

    doc.add_page_break()


def add_executive_summary(doc, report: FTOReport, section_title: str) -> None:
    """Add executive summary section."""
    from docx.shared import Pt

    doc.add_heading(section_title, level=1)

    governed_risk = governed_risk_level(report)
    risk = risk_display(governed_risk)
    risk_color = RISK_FILL.get(governed_risk, BRAND_SECONDARY_TEXT)

    p = doc.add_paragraph()
    run = p.add_run(f"Clearance Decision: {governed_decision_label(report)} ({risk} risk)")
    run.font.name = DOCX_FONTS["heading"]
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = _rgb_color(risk_color)

    add_styled_paragraph(
        doc,
        f"Blocking Patents: {governed_blocking_count(report)} | "
        f"Total Analyzed: {report.risk_summary.total_patents_analyzed} | "
        f"Total Discovered: {report.total_patents_found}",
        bold=True,
        size=10,
        color=BRAND_SECONDARY_TEXT,
    )

    doc.add_heading("Summary", level=2)
    add_styled_paragraph(doc, governed_executive_summary(report))


def add_evidence_scope(doc, report: FTOReport, section_title: str) -> None:
    """Add front-matter evidence scope before risk conclusions."""
    doc.add_heading(section_title, level=1)

    summary = summarize_evidence_scope(report)
    add_styled_paragraph(
        doc,
        f"{summary.headline}. {summary.posture}.",
        bold=True,
        color=BRAND_DEEP_TEAL,
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    style_header_cell(table.rows[0].cells[0], "Evidence Signal")
    style_header_cell(table.rows[0].cells[1], "Reader-Facing Scope")

    rows = [
        (
            "Successful sources",
            format_source_list(
                summary.successful_sources,
                empty="No successful-source telemetry recorded.",
            ),
        ),
        (
            "Unavailable sources",
            format_source_list(summary.unavailable_sources, empty="None recorded."),
        ),
        (
            "Skipped sources",
            format_source_list(summary.skipped_sources, empty="None recorded."),
        ),
        ("Confidence impact", summary.confidence_impact),
        ("Counsel review note", summary.review_note),
    ]

    for label, value in rows:
        row = table.add_row()
        style_body_cell(row.cells[0], label)
        style_body_cell(row.cells[1], value)

    if report.source_health.entries:
        doc.add_heading("Source Telemetry", level=2)
        health_table = doc.add_table(rows=1, cols=4)
        health_table.style = "Table Grid"
        style_header_cell(health_table.rows[0].cells[0], "Source")
        style_header_cell(health_table.rows[0].cells[1], "Status")
        style_header_cell(health_table.rows[0].cells[2], "Patents")
        style_header_cell(health_table.rows[0].cells[3], "Detail")

        for entry in report.source_health.entries:
            row = health_table.add_row()
            style_body_cell(row.cells[0], entry.source)
            label = source_status_label(entry)
            style_body_cell(row.cells[1], label)
            if entry.status.value in {"failed", "not_configured"}:
                set_cell_shading(row.cells[1], RISK_BG[get_risk_level("medium")])
            elif entry.status.value == "ok":
                set_cell_shading(row.cells[1], RISK_BG[get_risk_level("clear")])
            style_body_cell(row.cells[2], str(entry.patent_count))
            style_body_cell(row.cells[3], source_status_detail(entry))


def add_compound_profile(doc, report: FTOReport, section_title: str) -> None:
    """Add compound profile section with properties table."""
    doc.add_heading(section_title, level=1)

    c = report.compound
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    style_header_cell(table.rows[0].cells[0], "Property")
    style_header_cell(table.rows[0].cells[1], "Value")

    properties = [
        ("Name", c.name),
        ("SMILES", c.canonical_smiles),
        ("InChIKey", c.inchi_key),
        ("Molecular Formula", c.molecular_formula),
        ("Molecular Weight", f"{c.molecular_weight:.2f}" if c.molecular_weight else "N/A"),
    ]
    if c.cas_numbers:
        properties.append(("CAS Number(s)", ", ".join(c.cas_numbers)))
    if c.functional_groups:
        properties.append(("Functional Groups", ", ".join(c.functional_groups)))

    for prop_name, prop_value in properties:
        row = table.add_row()
        style_body_cell(row.cells[0], prop_name)
        style_body_cell(row.cells[1], str(prop_value))

    if c.related_compounds:
        doc.add_heading("Related Compounds", level=2)
        rel_table = doc.add_table(rows=1, cols=3)
        rel_table.style = "Table Grid"
        style_header_cell(rel_table.rows[0].cells[0], "Name")
        style_header_cell(rel_table.rows[0].cells[1], "SMILES")
        style_header_cell(rel_table.rows[0].cells[2], "Tanimoto")
        for rc in c.related_compounds[:10]:
            row = rel_table.add_row()
            style_body_cell(row.cells[0], rc.name)
            style_body_cell(row.cells[1], rc.canonical_smiles[:60])
            style_body_cell(row.cells[2], f"{rc.tanimoto_similarity:.3f}")


def add_methodology(doc, report: FTOReport, section_title: str) -> None:
    """Add search methodology section."""
    doc.add_heading(section_title, level=1)

    summary = summarize_evidence_scope(report)
    if report.source_health.entries:
        source_line = "Configured source requests: " + ", ".join(
            entry.source for entry in report.source_health.entries
        )
    else:
        source_line = "Recorded source names: " + (
            ", ".join(report.search_sources_used) or "Not recorded"
        )

    add_styled_paragraph(
        doc,
        f"Total patents discovered: {report.total_patents_found}\n"
        f"Patents after triage: {report.patents_after_triage}\n"
        f"Patents analyzed: {len(report.patent_analyses)}\n"
        f"Source status: {summary.headline}\n"
        f"{source_line}\n"
        f"Reader scope note: {summary.review_note}",
    )

    if report.source_health.entries:
        doc.add_heading("Source Health", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        style_header_cell(table.rows[0].cells[0], "Source")
        style_header_cell(table.rows[0].cells[1], "Status")
        style_header_cell(table.rows[0].cells[2], "Patents Found")

        for entry in report.source_health.entries:
            row = table.add_row()
            style_body_cell(row.cells[0], entry.source)
            status_text = source_status_label(entry)
            style_body_cell(row.cells[1], status_text)
            if entry.status.value == "failed":
                set_cell_shading(row.cells[1], RISK_BG[get_risk_level("high")])
            elif entry.status.value == "ok":
                set_cell_shading(row.cells[1], RISK_BG[get_risk_level("clear")])
            style_body_cell(row.cells[2], str(entry.patent_count))

    if report.audit_trail.timing_data:
        doc.add_heading("Pipeline Timing", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        style_header_cell(table.rows[0].cells[0], "Stage")
        style_header_cell(table.rows[0].cells[1], "Patents In")
        style_header_cell(table.rows[0].cells[2], "Patents Out")
        style_header_cell(table.rows[0].cells[3], "Duration")

        for step in report.audit_trail.timing_data:
            row = table.add_row()
            style_body_cell(row.cells[0], step.step_name)
            style_body_cell(row.cells[1], str(step.items_processed))
            style_body_cell(row.cells[2], str(step.items_output))
            style_body_cell(row.cells[3], f"{step.duration_seconds:.1f}s")


def add_risk_matrix(doc, report: FTOReport, section_title: str, *, options=None) -> None:
    """Add color-coded risk matrix table."""
    if not report.patent_analyses:
        return

    doc.add_heading(section_title, level=1)

    scientist_projection = getattr(options, "audience", "full") == "scientist"
    headers = ["Patent ID", "Title", "Assignee", "Risk", "Expiry"]
    if not scientist_projection:
        headers.extend(["Claims", "Orange Book"])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        style_header_cell(table.rows[0].cells[i], header)

    sorted_analyses = sorted(report.patent_analyses, key=lambda a: risk_sort_key(a.risk_level))

    for analysis in sorted_analyses:
        row = table.add_row()
        style_body_cell(row.cells[0], analysis.patent_id)
        style_body_cell(row.cells[1], analysis.title[:100])
        style_body_cell(row.cells[2], analysis.assignee)

        risk_text = risk_display(analysis.risk_level)
        style_body_cell(row.cells[3], risk_text)
        risk_bg = RISK_BG.get(analysis.risk_level, BRAND_PAPER)
        set_cell_shading(row.cells[3], risk_bg)

        style_body_cell(
            row.cells[4],
            analysis.expiry_date.isoformat() if analysis.expiry_date else "N/A",
        )
        if not scientist_projection:
            claims_str = ", ".join(str(c.claim_number) for c in analysis.claims_analyzed[:10])
            style_body_cell(row.cells[5], claims_str)

            ob_text = ""
            if analysis.orange_book_info and analysis.orange_book_info.is_listed:
                ob_text = (
                    "LISTED — DELIST REQUESTED"
                    if analysis.orange_book_info.delist_requested
                    else "LISTED"
                )
            style_body_cell(row.cells[6], ob_text)
