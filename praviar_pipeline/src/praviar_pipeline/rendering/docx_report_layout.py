"""Shared layout helpers for DOCX report rendering."""

from __future__ import annotations

from typing import TYPE_CHECKING

from praviar_pipeline.rendering.brand_mark import render_praviar_mark_png_stream
from praviar_pipeline.rendering.branding import (
    SUPPORTED_OFFICE_BRANDING_LOGO_EXTENSIONS,
    resolve_branding_logo_path,
)
from praviar_pipeline.rendering.design import (
    BRAND_BODY_TEXT,
    BRAND_INK,
    BRAND_PAPER,
    BRAND_SECONDARY_TEXT,
    DOCX_FONTS,
    DOCX_SIZES,
    hex_to_rgb,
)

if TYPE_CHECKING:
    from praviar_pipeline.rendering.branding import BrandingConfig


def _add_field(paragraph, instruction: str) -> None:
    """Add a Word field such as PAGE or NUMPAGES to a paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    run._r.append(instr_text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def rgb_color(hex_str: str):
    """Create a python-docx RGBColor from a hex string."""
    from docx.shared import RGBColor

    r, g, b = hex_to_rgb(hex_str)
    return RGBColor(r, g, b)


def set_cell_shading(cell, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    from docx.oxml.ns import qn
    from lxml import etree

    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), hex_color.lstrip("#"))
    shading.set(qn("w:val"), "clear")


def add_styled_paragraph(
    doc,
    text: str,
    font: str = DOCX_FONTS["body"],
    size: float = DOCX_SIZES["body"],
    bold: bool = False,
    italic: bool = False,
    color: str = BRAND_BODY_TEXT,
    space_after_pt: float = 6,
) -> None:
    """Add a paragraph with full styling."""
    from docx.shared import Pt

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb_color(color)
    p.paragraph_format.space_after = Pt(space_after_pt)


def style_header_cell(cell, text: str) -> None:
    """Style a table header cell with ink background and paper text."""
    from docx.shared import Pt

    cell.text = text
    set_cell_shading(cell, BRAND_INK)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = DOCX_FONTS["table"]
            run.font.size = Pt(DOCX_SIZES["table_header"])
            run.font.bold = True
            run.font.color.rgb = rgb_color(BRAND_PAPER)


def style_body_cell(cell, text: str, font_size: float = DOCX_SIZES["table_body"]) -> None:
    """Style a table body cell."""
    from docx.shared import Pt

    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = DOCX_FONTS["table"]
            run.font.size = Pt(font_size)
            run.font.color.rgb = rgb_color(BRAND_BODY_TEXT)


def status_bg(status: str) -> str:
    """Map element status to background color."""
    return {
        "met": "#FDECEC",
        "partially_met": "#F7EEE5",
        "not_met": "#D7ECE5",
        "unclear": BRAND_PAPER,
    }.get(status, BRAND_PAPER)


def get_risk_level(value: str):
    """Get RiskLevel enum from string, importing lazily."""
    from praviar_pipeline.models.analysis import RiskLevel

    mapping = {
        "high": RiskLevel.HIGH,
        "medium": RiskLevel.MEDIUM,
        "low": RiskLevel.LOW,
        "clear": RiskLevel.CLEAR,
    }
    return mapping.get(value.lower(), RiskLevel.MEDIUM)


def configure_document(doc, branding: BrandingConfig) -> None:
    """Apply document-wide fonts and header/footer branding."""
    from docx.shared import Inches, Pt

    style = doc.styles["Normal"]
    style.font.name = DOCX_FONTS["body"]
    style.font.size = Pt(DOCX_SIZES["body"])
    style.font.color.rgb = rgb_color(BRAND_BODY_TEXT)

    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False

        section.header.is_linked_to_previous = False
        header_para = section.header.paragraphs[0]
        logo_path = resolve_branding_logo_path(
            branding,
            renderer_name="DOCX",
            supported_extensions=SUPPORTED_OFFICE_BRANDING_LOGO_EXTENSIONS,
        )
        if logo_path is not None:
            header_para.add_run().add_picture(str(logo_path), width=Inches(0.28))
        elif not branding.suppresses_praviar_branding:
            header_para.add_run().add_picture(
                render_praviar_mark_png_stream(variant="on_light", size_px=96),
                width=Inches(0.24),
            )
        header_run = header_para.add_run(branding.header_text)
        header_run.font.name = DOCX_FONTS["table"]
        header_run.font.size = Pt(8)
        header_run.font.color.rgb = rgb_color(BRAND_SECONDARY_TEXT)

        section.footer.is_linked_to_previous = False
        footer_para = section.footer.paragraphs[0]
        footer_run = footer_para.add_run(branding.footer_text)
        footer_run.font.name = DOCX_FONTS["table"]
        footer_run.font.size = Pt(7)
        footer_run.font.color.rgb = rgb_color(BRAND_SECONDARY_TEXT)

        page_prefix = footer_para.add_run(" | Page ")
        page_prefix.font.name = DOCX_FONTS["table"]
        page_prefix.font.size = Pt(7)
        page_prefix.font.color.rgb = rgb_color(BRAND_SECONDARY_TEXT)
        _add_field(footer_para, "PAGE")

        page_total = footer_para.add_run(" of ")
        page_total.font.name = DOCX_FONTS["table"]
        page_total.font.size = Pt(7)
        page_total.font.color.rgb = rgb_color(BRAND_SECONDARY_TEXT)
        _add_field(footer_para, "NUMPAGES")
